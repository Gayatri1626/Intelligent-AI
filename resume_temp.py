from docx import Document

# Define a function to generate a resume
def generate_resume(template, name, email, phone, education, experience, skills, job_title=None, job_description=None, objective=None):
    # Initialize a new Document
    doc = Document()
    
    # Add a title based on the template or job title
    if job_title:
        doc.add_heading(f"{name} - {job_title}", level=1)
    else:
        doc.add_heading(name, level=1)
    
    # Add job description if provided
    if job_description:
        doc.add_heading("Job Description", level=2)
        doc.add_paragraph(job_description)
        doc.add_paragraph()
    
    # Add objective if provided
    if objective:
        doc.add_heading("Objective", level=2)
        doc.add_paragraph(objective)
        doc.add_paragraph()
    
    # Add contact information
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    doc.add_paragraph()
    
    # Add education section
    doc.add_heading("Education", level=2)
    doc.add_paragraph(education)
    doc.add_paragraph()
    
    # Add experience section
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(experience)
    doc.add_paragraph()
    
    # Add skills section
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(skills)
    
    # Save the document with a specific filename
    file_path = f"{name.replace(' ', '_')}_Resume.docx"
    doc.save(file_path)
    
    return file_path
